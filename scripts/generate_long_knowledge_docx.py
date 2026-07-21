from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("docs/enterprise_service_knowledge_base_longform.docx")


def set_east_asia_font(style, font_name: str) -> None:
    style.font.name = font_name
    rpr = style.element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.append(fonts)
    fonts.set(qn("w:eastAsia"), font_name)


def add_paragraph(doc: Document, text: str, style: str = "Normal") -> None:
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_heading(level=level)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(46, 116, 181)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(46, 116, 181)
    else:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(31, 77, 120)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3", "List Bullet"]:
        style = doc.styles[style_name]
        set_east_asia_font(style, "PingFang SC")
        if style_name == "Normal":
            style.font.name = "Arial"
            style.font.size = Pt(11)
            style.paragraph_format.space_after = Pt(6)
            style.paragraph_format.line_spacing = 1.25


CHAPTERS = [
    (
        "第一章 企業服務定位與溝通原則",
        "企業服務定位",
        "品牌承諾、服務語氣和解釋邊界",
        ["企業能提供什麼服務", "客服身份說明", "處理依據解釋"],
    ),
    (
        "第二章 用戶身份、稱呼與隱私保護",
        "身份與隱私",
        "用戶稱呼、賬號標識和敏感信息保護",
        ["用戶更改稱呼", "用戶提供會員賬號", "用戶要求使用歷史地址"],
    ),
    (
        "第三章 商品信息、價格解釋與購買諮詢",
        "商品與價格",
        "價格來源、商品規格和購買意向確認",
        ["商品價格諮詢", "商品規格對比", "購買前確認"],
    ),
    (
        "第四章 訂單創建、支付確認與取消處理",
        "訂單處理",
        "下單確認、支付狀態和取消窗口",
        ["創建訂單", "支付核對", "取消剛創建的訂單"],
    ),
    (
        "第五章 物流履約、配送承諾與改派協同",
        "物流履約",
        "發貨進度、配送承諾和地址改派",
        ["訂單未發貨", "配送晚到", "用戶要求改地址"],
    ),
    (
        "第六章 退款、退貨與換貨處理",
        "售後處理",
        "退款條件、退貨回收和換貨履約",
        ["申請退款", "申請退貨", "換顏色或換規格"],
    ),
    (
        "第七章 會員等級、權益發放與補償",
        "會員權益",
        "等級資格、權益發放和補償邊界",
        ["會員券未到賬", "積分缺失", "活動禮品延遲"],
    ),
    (
        "第八章 投訴、風險與人工介入",
        "投訴風險",
        "升級路徑、證據保全和響應時限",
        ["用戶投訴", "要求人工", "涉及隱私或扣款風險"],
    ),
    (
        "第九章 企業文化、品牌語氣與服務紅線",
        "企業文化",
        "服務語氣、禁用表達和承諾邊界",
        ["用戶質疑態度", "要求保證結果", "要求內部規則"],
    ),
    (
        "第十章 內部運營排查與數據口徑",
        "運營排查",
        "指標口徑、工單歸因和復盤字段",
        ["差評上升", "責任歸因", "處理耗時復盤"],
    ),
    (
        "第十一章 知識維護、工具協作與流程治理",
        "流程治理",
        "知識維護、工具調用和流程更新",
        ["文檔更新", "接口能力變化", "流程發佈回滾"],
    ),
    (
        "第十二章 服務質量復盤與持續改進",
        "質量改進",
        "質量檢查、問題復盤和改進閉環",
        ["處理結果復盤", "服務標準調整", "跨團隊協作改進"],
    ),
]


def chapter_paragraphs(chapter: str, domain: str, concern: str, examples: list[str]) -> list[str]:
    example_text = "、".join(examples)
    return [
        (
            f"{chapter}用於統一{domain}相關事項的處理方式。服務人員接到{example_text}等問題時，"
            f"應先確認用戶真實訴求、當前已知事實、可執行動作和需要補充的信息。若用戶表達中同時包含多個事項，"
            f"應區分主訴求與後續訴求，先處理時效更強、風險更高或用戶明確要求優先處理的部分。"
        ),
        (
            f"{domain}場景下的回覆需要同時滿足準確、清楚和可執行三項要求。準確是指結論必須來自已確認事實、"
            f"已發佈規則或可追溯記錄；清楚是指用戶能理解當前狀態和下一步；可執行是指回復後能形成追問、查詢、"
            f"創建、取消、補償、轉人工或結束等明確動作。"
        ),
        (
            f"涉及{concern}時，服務人員應避免過度承諾。能夠立即辦理的事項，應說明已辦理結果和後續影響；"
            f"需要等待外部處理的事項，應說明預計觀察點、責任團隊和用戶可以採取的下一步；需要補充信息的事項，"
            f"應一次性列出關鍵缺口，避免連續多輪只追問一個字段。"
        ),
        (
            f"如果同一問題存在多個規則來源，應優先採用發佈時間較新、適用範圍較窄、與用戶條件更匹配的內容。"
            f"當規則之間存在明顯衝突時，不應自行擴大解釋，而應保留衝突點並轉給相應負責人確認。"
            f"對用戶側表達，應使用結論先行的語言，內部原因只在必要時簡要說明。"
        ),
        (
            f"{domain}事項處理完成後，應留下能夠復盤的摘要，包括用戶訴求、關鍵事實、採取動作、未解決風險和後續責任。"
            f"若用戶繼續提出新事項，應在原事項完成狀態明確後再進入下一事項，避免把不同事項的字段、工具結果或處理結論混在一起。"
        ),
    ]


def build_document() -> Document:
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title_run = title.add_run("面壁智能客戶服務知識手冊")
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(11, 37, 69)
    title.paragraph_format.space_after = Pt(4)

    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("企業介紹、服務原則、訂單履約、會員權益與售後處理參考")
    subtitle_run.font.name = "Arial"
    subtitle_run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = RGBColor(85, 85, 85)

    add_paragraph(
        doc,
        "版本：2026-06-16。本文檔面向客戶服務、運營管理、產品支持和服務質量團隊，"
        "彙總企業客戶服務中的常見規則、處理原則、溝通邊界和跨團隊協作方式。"
        "文檔以章節和段落組織，適合長期維護、定期復盤和按業務主題擴展。",
    )
    add_heading(doc, "文檔說明", 1)
    for item in [
        "本手冊適用於客戶服務、運營復盤、流程治理和知識庫維護。",
        "處理用戶問題時，應優先區分事實核對、政策解釋、執行動作和風險升級。",
        "涉及敏感信息、補償承諾或人工介入時，應保留必要依據並避免過度承諾。",
        "當用戶一次提出多個訴求時，應先明確優先級，再按事項形成獨立處理結論。",
    ]:
        add_bullet(doc, item)

    for chapter_index, (chapter, domain, concern, examples) in enumerate(CHAPTERS, start=1):
        add_heading(doc, chapter, 1)
        add_paragraph(
            doc,
            f"本章圍繞{domain}展開，關注{concern}。相關內容適用於日常客服對話、內部工單流轉、"
            "運營排查和服務質量復盤。服務人員應根據用戶當前訴求和已確認事實選擇合適處理方式。",
        )
        for section_index in range(1, 5):
            add_heading(doc, f"{chapter_index}.{section_index} {domain}處理原則 {section_index}", 2)
            for paragraph in chapter_paragraphs(chapter, domain, concern, examples):
                add_paragraph(
                    doc,
                    paragraph
                    + f"在第 {chapter_index} 章第 {section_index} 節的具體執行中，還應結合當前業務狀態、"
                    f"用戶歷史偏好、可用工具返回、知識庫記錄和人工審核要求，形成清晰的處理鏈路。"
                )
            add_heading(doc, f"{chapter_index}.{section_index}.1 常見場景說明", 3)
            for item in [
                f"當用戶直接描述{examples[0]}時，應先確認是否已有可執行信息，再決定查詢、辦理或解釋。",
                f"當用戶同時涉及{examples[1]}和其他事項時，應拆清事項邊界，避免把不同處理結果混用。",
                f"當用戶表達{examples[2]}時，應說明當前可做動作、限制條件和下一步責任歸屬。",
            ]:
                add_bullet(doc, item)

    add_heading(doc, "跨章節協作原則", 1)
    cross_domain_topics = [
        (
            "會員權益與訂單售後經常同時出現。例如用戶反饋會員券未到賬又要求取消訂單時，"
            "應先確認訂單狀態，再核對權益資格和發放記錄，最後說明取消訂單對權益的影響。"
        ),
        (
            "物流履約與補償承諾也經常交叉。用戶認為配送晚到時，需要確認承諾來源、發貨狀態、"
            "實際簽收時間和補償規則，不應只憑用戶口述立即承諾補償。"
        ),
        (
            "投訴升級與人工介入需要保留完整上下文。轉交前應整理用戶訴求、已核對字段、已嘗試動作、"
            "未解決風險和建議處理方向，減少用戶重複描述。"
        ),
        (
            "服務質量復盤應同時查看對話過程和業務結果。單次負面反饋只能說明用戶在該輪體驗不滿意，"
            "還需要結合回覆是否準確、流程是否完整、工具是否可用和最終問題是否解決。"
        ),
    ]
    for repeat in range(16):
        for topic in cross_domain_topics:
            add_paragraph(
                doc,
                f"跨章節協作說明 {repeat + 1}：{topic}"
                "處理跨域事項時，應保持事項獨立、證據清楚、結論明確，並在用戶可理解的範圍內說明下一步。"
            )

    add_heading(doc, "服務知識維護原則", 1)
    maintenance_topics = [
        "知識內容應按業務負責人、更新時間、適用範圍和例外條件進行維護。過期內容應及時下線，但歷史版本仍需保留用於審計。",
        "工具能力發生變化時，應同步更新相關流程說明、可執行動作和異常處理方式，避免前端說明與後端能力不一致。",
        "不同智能體可以擁有不同可見範圍。分支中的知識或技能改動不應影響整體版本，除非經過負責人確認並推送到整體。",
        "新規則發佈前應檢查是否影響現有售後、會員、物流和訂單流程，尤其要關注金額、承諾、隱私和人工介入邊界。",
    ]
    for repeat in range(12):
        for topic in maintenance_topics:
            add_paragraph(
                doc,
                f"維護說明 {repeat + 1}：{topic}"
                "維護動作完成後，應記錄變更原因、影響範圍和回滾方式，便於後續復盤。"
            )

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT)
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    cjk_count = len(re.findall(r"[\u4e00-\u9fff]", text))
    print(f"written={OUTPUT}")
    print(f"paragraphs={len(doc.paragraphs)}")
    print(f"tables={len(doc.tables)}")
    print(f"cjk_chars={cjk_count}")


if __name__ == "__main__":
    main()
